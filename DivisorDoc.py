from docx import Document
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError
from io import BytesIO
from PIL import Image
from PIL import UnidentifiedImageError
import os
import re

def get_or_create_style(new_doc, template_doc, style_name):
    try:
        return new_doc.styles[style_name]
    except KeyError:
        try:
            template_style = template_doc.styles[style_name]
            new_style = new_doc.styles.add_style(style_name, template_style.type)
            new_style.base_style = template_style
            return new_style
        except KeyError:
            new_style = new_doc.styles.add_style(style_name, 1)
            return new_style

def extract_images(run, new_run):
    for blip in run.element.xpath('.//a:blip'):
        rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        image_part = run.part.related_parts[rId]

        try:
            # Tentar abrir a imagem
            image = Image.open(BytesIO(image_part.blob))
            width, height = image.size
            dpi = image.info.get('dpi', (96, 96))[0]  # Assume 96 DPI se não especificado
            image.close()

            # Converter o tamanho da imagem para polegadas para docx
            width_in_inches = width / dpi
            height_in_inches = height / dpi

            # Adicionar imagem ao novo documento com tamanho original
            new_run.add_picture(BytesIO(image_part.blob), width=Inches(width_in_inches), height=Inches(height_in_inches))

        except (UnrecognizedImageError, UnidentifiedImageError):
            print(f"Erro ao reconhecer a imagem: {rId}")
            continue
        except Exception as e:
            print(f"Erro ao processar imagem {rId}: {e}")
            continue
        

def copy_paragraph(paragraph, new_doc, template_doc):
    new_paragraph = new_doc.add_paragraph()

    # Buscar e aplicar o estilo ao novo parágrafo
    style_name = paragraph.style.name
    new_style = get_or_create_style(new_doc, template_doc, style_name)
    new_paragraph.style = new_style

    # Copiar formatação do parágrafo
    new_paragraph.alignment = paragraph.alignment
    new_paragraph.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
    new_paragraph.paragraph_format.right_indent = paragraph.paragraph_format.right_indent
    new_paragraph.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
    new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
    new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after
    new_paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
    new_paragraph.paragraph_format.keep_together = paragraph.paragraph_format.keep_together
    new_paragraph.paragraph_format.keep_with_next = paragraph.paragraph_format.keep_with_next
    new_paragraph.paragraph_format.page_break_before = paragraph.paragraph_format.page_break_before
    new_paragraph.paragraph_format.widow_control = paragraph.paragraph_format.widow_control

    # Identificar se o parágrafo é uma lista e obter detalhes da numeração
    numPr = paragraph._element.xpath('.//w:numPr')
    if numPr:
        numId = numPr[0].xpath('.//w:numId/@w:val')[0]
        ilvl = numPr[0].xpath('.//w:ilvl/@w:val')[0]
        new_paragraph._element.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = int(numId)
        new_paragraph._element.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = int(ilvl)

    # Copiar runs
    for run in paragraph.runs:
        new_run = new_paragraph.add_run(run.text)
        
        # Copiar estilo de run
        new_run.font.name = run.font.name
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.color.rgb = run.font.color.rgb

        extract_images(run, new_run)

    return new_paragraph

def split_document_by_heading(doc_path, output_dir, template_path):
    # Abrir o documento original
    doc = Document(doc_path)

    # Carregar o template para estilos
    template_doc = Document(template_path)

    # Cria a pasta de saída se não existir
    os.makedirs(output_dir, exist_ok=True)

    current_doc = None
    current_heading = None
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 1':
            if current_doc:
                # Salvar o documento atual
                current_doc.save(os.path.join(output_dir, f"{re.sub(r'[\\/*?:\"<>|]', '', current_heading)}.docx"))

            # Iniciar um novo documento baseado no template
            current_doc = Document(template_path)
            current_heading = paragraph.text

            # Adicionar o título ao novo documento
            copy_paragraph(paragraph, current_doc, template_doc)
        elif paragraph.style.name == 'Heading 2':
            # Adicionar subtítulos com estilo de Heading 2
            copy_paragraph(paragraph, current_doc, template_doc)
        elif current_doc:
            # Copiar parágrafos e outros elementos ao documento atual
            copy_paragraph(paragraph, current_doc, template_doc)

    # Salvar o último documento
    if current_doc:
        current_doc.save(os.path.join(output_dir, f"{re.sub(r'[\\/*?:\"<>|]', '', current_heading)}.docx"))

doc_path = ""
output_dir = ""
template_path = ""
split_document_by_heading(doc_path, output_dir, template_path)
